import os
import json
from google import genai
from pydantic import BaseModel
from flask import Flask, request, jsonify, send_from_directory, render_template, session
from werkzeug.utils import secure_filename
from docx import Document
import pypdf

# --------------------------------------
# Configuration
# --------------------------------------
# os.environ['GENAI_API_KEY'] = 'YOUR_API_KEY_HERE' # Set this in your environment variables
UPLOAD_FOLDER = 'UPLOAD_FOLDER'
ALLOWED_EXTENSIONS = {'pdf', 'docx'}
GEMINI_API_KEY = os.environ.get("GENAI_API_KEY")

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB upload limit
app.secret_key = 'super_secret_key_for_demo_only' # In prod, use a real secret key

client = None
if GEMINI_API_KEY:
    client = genai.Client(api_key=GEMINI_API_KEY)

# In-memory storage for demo purposes (use a DB in production)
# Structure: { session_id: { 'text': ..., 'variables': [...], 'answers': {...}, 'filename': ... } }
user_sessions = {}

# --------------------------------------
# Models
# --------------------------------------
class Variable(BaseModel):
    name: str
    description: str

class VariableList(BaseModel):
    variables: list[Variable]

# --------------------------------------
# Helper
# --------------------------------------
def allowed_file(filename: str) -> bool:
    """Check if file extension is allowed."""
    return (
        '.' in filename 
        and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS
    )

def extract_text_from_docx(filepath):
    doc = Document(filepath)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def extract_text_from_pdf(filepath):
    text = ""
    with open(filepath, 'rb') as file:
        reader = pypdf.PdfReader(file)
        for page in reader.pages:
            text += page.extract_text() + "\n"
    return text

def identify_variables_with_gemini(text):
    if not client:
        print("DEBUG: Using mock data because client is not initialized.")
        # Mock response if no API key or client not initialized
        return [
            {"name": "Client Name", "description": "The full name of the client"},
            {"name": "Date", "description": "The date of the agreement"},
            {"name": "Amount", "description": "The total amount in USD"}
        ]
    
    prompt = f"""
    Analyze the following legal document text and identify all the variable fields that need to be filled in by the user.
    Ignore standard boilerplate text. Look for placeholders like [Name], {{Date}}, or contextually missing information.
    
    Document Text:
    {text[:10000]} # Truncate to avoid token limits for this demo
    """
    
    try:
        response = client.models.generate_content(
            model='gemini-2.0-flash',
            contents=prompt,
            config={
                'response_mime_type': 'application/json',
                'response_schema': VariableList,
            },
        )
        # response.parsed is already a VariableList instance
        return [v.model_dump() for v in response.parsed.variables]
    except Exception as e:
        print(f"Gemini Error: {e}")
        return []

def replace_variables_in_docx(filepath, answers, output_path):
    doc = Document(filepath)
    # Simple replacement in paragraphs
    # Note: This is a basic implementation. Complex docx structures (tables, headers) might need more work.
    for para in doc.paragraphs:
        for key, value in answers.items():
            # This is a naive replacement. Real-world docx replacement is complex due to runs.
            # For this demo, we'll assume simple text matching or placeholders.
            # Since we don't know the exact placeholder format from Gemini, we might need a smarter way.
            # For now, let's assume we append the values at the end or try to find the "name" if it exists.
            # BETTER APPROACH for this demo: Just append a summary of filled variables at the end
            # because exact in-place replacement requires knowing the EXACT placeholder string used in the doc.
            pass
            
    # For the sake of the demo "filling placeholders", we will create a new document 
    # that lists the filled terms, or try to replace if we can guess the placeholder.
    # Let's try to replace standard bracketed placeholders if they match the variable name.
    
    for para in doc.paragraphs:
        for key, value in answers.items():
            if key in para.text:
                para.text = para.text.replace(key, value)
            # Also try brackets
            if f"[{key}]" in para.text:
                para.text = para.text.replace(f"[{key}]", value)
                
    doc.save(output_path)

# --------------------------------------
# Routes
# --------------------------------------
@app.route('/')
def index():
    """Serve frontend home page."""
    return render_template('index.html')


@app.route('/upload', methods=['POST'])
def upload_file():
    """Handle file upload."""
    if 'file' not in request.files:
        return jsonify({"status": "error", "message": "No file part"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({"status": "error", "message": "No file selected"}), 400

    if not allowed_file(file.filename):
        return jsonify({"status": "error", "message": "Invalid file type"}), 400

    try:
        filename = secure_filename(file.filename)
        save_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
        file.save(save_path)
        
        # Initialize session
        session_id = os.urandom(16).hex()
        user_sessions[session_id] = {
            'filename': filename,
            'filepath': save_path,
            'answers': {}
        }
        
        return jsonify({
            "status": "success",
            "message": "File uploaded.",
            "session_id": session_id
        })

    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route('/analyze', methods=['POST'])
def analyze_document():
    data = request.json
    session_id = data.get('session_id')
    
    if not session_id or session_id not in user_sessions:
        return jsonify({"status": "error", "message": "Invalid session"}), 400
        
    session_data = user_sessions[session_id]
    filepath = session_data['filepath']
    
    # Extract text
    try:
        if filepath.endswith('.pdf'):
            text = extract_text_from_pdf(filepath)
        else:
            text = extract_text_from_docx(filepath)
            
        session_data['text'] = text
        
        # Identify variables
        variables = identify_variables_with_gemini(text)
        session_data['variables'] = variables
        
        return jsonify({
            "status": "success",
            "variables": variables
        })
        
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route('/generate', methods=['POST'])
def generate_document():
    data = request.json
    session_id = data.get('session_id')
    answers = data.get('answers') # { "Variable Name": "Value" }
    
    if not session_id or session_id not in user_sessions:
        return jsonify({"status": "error", "message": "Invalid session"}), 400
        
    session_data = user_sessions[session_id]
    original_path = session_data['filepath']
    
    # Create output filename
    filename_base = os.path.splitext(session_data['filename'])[0]
    output_filename = f"{filename_base}_filled.docx"
    output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
    
    try:
        # For PDF we can't easily edit, so we'll just create a summary DOCX for now
        # Or if it was a DOCX, we try to replace.
        if original_path.endswith('.docx'):
            replace_variables_in_docx(original_path, answers, output_path)
        else:
            # Create a new doc with the answers
            doc = Document()
            doc.add_heading('Filled Variables', 0)
            for k, v in answers.items():
                doc.add_paragraph(f"{k}: {v}")
            doc.save(output_path)
            
        return jsonify({
            "status": "success",
            "download_url": f"/download/{output_filename}"
        })
        
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


@app.route('/download/<filename>', methods=['GET'])
def download_file(filename):
    """Serve file securely to client."""
    try:
        return send_from_directory(
            app.config['UPLOAD_FOLDER'],
            filename,
            as_attachment=True
        )
    except Exception:
        return jsonify({"status": "error", "message": "File not found"}), 404


if __name__ == "__main__":
    app.run(debug=True, port=5001)
